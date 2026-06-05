from __future__ import annotations

from datetime import date, datetime
from html.parser import HTMLParser
import json
import os
import sys
import warnings
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".xlsx",
    ".xls",
    ".csv",
    ".tsv",
    ".txt",
    ".md",
    ".markdown",
    ".rst",
    ".html",
    ".htm",
}

TEXT_EXTENSIONS = {".txt", ".md", ".markdown", ".rst"}
HTML_EXTENSIONS = {".html", ".htm"}
SPREADSHEET_EXTENSIONS = {".xlsx", ".xls", ".csv", ".tsv"}
SKIP_DIRS = {
    "node_modules",
    ".git",
    ".svn",
    ".hg",
    ".idea",
    ".vscode",
    ".vs",
    "dist",
    "build",
    "out",
    "bin",
    "obj",
    "target",
    "vendor",
    "__pycache__",
    ".next",
    ".nuxt",
    ".cache",
    ".turbo",
    "coverage",
    ".gradle",
    ".dart_tool",
    ".fvm",
}


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [line.rstrip() for line in text.split("\n")]
    compacted: list[str] = []
    blank_seen = False
    for line in lines:
        if line.strip():
            compacted.append(line.strip())
            blank_seen = False
        elif compacted and not blank_seen:
            compacted.append("")
            blank_seen = True
    return "\n".join(compacted).strip()


class HtmlTextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self._skip_depth = 0
        self._parts: list[str] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        tag = tag.lower()
        if tag in {"script", "style", "noscript", "svg"}:
            self._skip_depth += 1
            return
        if tag in {"p", "br", "div", "section", "article", "li", "tr", "h1", "h2", "h3", "h4", "h5", "h6"}:
            self._parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in {"script", "style", "noscript", "svg"} and self._skip_depth > 0:
            self._skip_depth -= 1
            return
        if tag in {"p", "div", "section", "article", "li", "tr", "h1", "h2", "h3", "h4", "h5", "h6"}:
            self._parts.append("\n")

    def handle_data(self, data: str) -> None:
        if self._skip_depth > 0:
            return
        text = data.strip()
        if text:
            self._parts.append(text)
            self._parts.append(" ")

    def text(self) -> str:
        return normalize_text("".join(self._parts))


def stringify_value(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and value != value:
        return ""
    if isinstance(value, bool):
        return "TRUE" if value else "FALSE"
    return str(value).strip()


def parse_optional_positive_int(value: Any) -> int | None:
    if value in (None, "", 0, "0"):
        return None
    parsed = int(value)
    return parsed if parsed > 0 else None


def parse_optional_non_negative_int(value: Any, default: int = 0) -> int:
    if value in (None, ""):
        return default
    parsed = int(value)
    return parsed if parsed >= 0 else default


def parse_focus_columns(raw: Any) -> list[str]:
    if not isinstance(raw, str):
        return []
    return [item.strip() for item in raw.split(",") if item.strip()]


def parse_comma_separated(raw: Any) -> list[str]:
    if raw in (None, ""):
        return []
    if isinstance(raw, str):
        return [item.strip() for item in raw.split(",") if item.strip()]
    if isinstance(raw, list):
        values: list[str] = []
        for item in raw:
            text = stringify_value(item)
            if text:
                values.append(text)
        return values
    return []


def parse_json_items(raw: Any, label: str) -> list[dict[str, Any]]:
    if raw in (None, "", []):
        return []

    parsed = raw
    if isinstance(raw, str):
        try:
            parsed = json.loads(raw)
        except json.JSONDecodeError as exc:
            raise RuntimeError(f"{label} must be a JSON array string: {exc.msg}") from exc

    if isinstance(parsed, dict):
        parsed = [parsed]

    if not isinstance(parsed, list):
        raise RuntimeError(f"{label} must be a JSON array.")

    items: list[dict[str, Any]] = []
    for index, item in enumerate(parsed, start=1):
        if not isinstance(item, dict):
            raise RuntimeError(f"{label}[{index}] must be an object.")
        items.append(item)
    return items


def relative_path(path: Path, workspace_root: Path) -> str:
    try:
        return path.relative_to(workspace_root).as_posix()
    except ValueError:
        return path.as_posix()


def parse_extensions(raw: str | None) -> set[str]:
    if not raw:
        return set(SUPPORTED_EXTENSIONS)
    parsed = {
        ("." + item.strip().lstrip(".").lower())
        for item in raw.split(",")
        if item.strip()
    }
    return parsed or set(SUPPORTED_EXTENSIONS)


def apply_limits(
    blocks: list[dict[str, Any]],
    max_chars: int,
    max_blocks: int,
) -> tuple[list[dict[str, Any]], str, int, bool]:
    full_text = "\n\n".join(block["text"] for block in blocks if block.get("text"))
    total_char_count = len(full_text)

    limited: list[dict[str, Any]] = []
    remaining = max_chars
    truncated = False

    for block in blocks:
        raw_text = block.get("text", "")
        if not raw_text:
            continue

        if len(limited) >= max_blocks or remaining <= 0:
            truncated = True
            break

        block_text = raw_text
        block_truncated = False
        if len(block_text) > remaining:
            keep = max(remaining - 3, 0)
            block_text = f"{block_text[:keep]}..." if keep > 0 else raw_text[:remaining]
            block_truncated = True
            truncated = True

        shaped = {
            key: value
            for key, value in block.items()
            if key != "text" and value is not None
        }
        shaped["text"] = block_text
        shaped["charCount"] = len(raw_text)
        if block_truncated:
            shaped["truncated"] = True
        limited.append(shaped)
        remaining -= len(block_text)

    if len(limited) < len([b for b in blocks if b.get("text")]):
        truncated = True

    content = "\n\n".join(block["text"] for block in limited if block.get("text"))
    return limited, content, total_char_count, truncated


def build_result(
    *,
    path: Path,
    workspace_root: Path,
    document_type: str,
    title: str | None,
    metadata: dict[str, Any],
    blocks: list[dict[str, Any]],
    max_chars: int,
    max_blocks: int,
) -> dict[str, Any]:
    limited_blocks, content, char_count, truncated = apply_limits(blocks, max_chars, max_blocks)
    metadata = dict(metadata)
    metadata["blockCount"] = len(blocks)
    if document_type == "pdf" and not content:
        metadata["warning"] = "No text extracted. This PDF may be image-based and require OCR."
    return {
        "path": relative_path(path, workspace_root),
        "documentType": document_type,
        "title": title,
        "content": content,
        "charCount": char_count,
        "truncated": truncated,
        "metadata": metadata,
        "blocks": limited_blocks,
    }


def load_tabular_frame(
    path: Path,
    *,
    sheet: str | None = None,
):
    try:
        import pandas as pd  # pylint: disable=import-outside-toplevel
    except Exception as exc:  # pragma: no cover
        raise RuntimeError(f"pandas is required to read {path.suffix} files: {exc}") from exc

    suffix = path.suffix.lower()
    metadata: dict[str, Any] = {}

    if suffix in {".csv", ".tsv"}:
        separator = "\t" if suffix == ".tsv" else ","
        frame = pd.read_csv(
            str(path),
            sep=separator,
            keep_default_na=True,
            low_memory=False,
        )
        frame.columns = [stringify_value(column) for column in frame.columns.tolist()]
        metadata["sheetName"] = None
        return frame, path.stem, metadata

    if suffix == ".xlsx":
        workbook = pd.ExcelFile(str(path))
        if not workbook.sheet_names:
            raise RuntimeError("Workbook does not contain any sheets.")
        selected_sheet = sheet or workbook.sheet_names[0]
        if selected_sheet not in workbook.sheet_names:
            raise RuntimeError(
                f"Sheet '{selected_sheet}' was not found. Available sheets: {', '.join(workbook.sheet_names)}"
            )
        frame = workbook.parse(sheet_name=selected_sheet)
        frame.columns = [stringify_value(column) for column in frame.columns.tolist()]
        metadata["sheetName"] = selected_sheet
        metadata["sheetNames"] = workbook.sheet_names
        return frame, selected_sheet, metadata

    if suffix == ".xls":
        raise RuntimeError(
            "Legacy .xls analysis is not supported by the current runtime. "
            "Please convert the file to .xlsx for reliable tabular analysis."
        )

    raise RuntimeError(f"Unsupported tabular file type: {suffix or '[no extension]'}")


def frame_window_to_blocks(
    source_name: str,
    frame,
    *,
    row_offset: int = 0,
    max_rows: int | None = None,
    chunk_size: int = 40,
) -> tuple[list[dict[str, Any]], int, int]:
    total_rows = len(frame.index)
    bounded_offset = max(0, min(row_offset, total_rows))
    window = frame.iloc[bounded_offset:]
    if max_rows is not None:
        window = window.iloc[:max_rows]

    header = [stringify_value(column) for column in window.columns.tolist()]
    window = window.fillna("")
    blocks: list[dict[str, Any]] = []

    for start in range(0, len(window.index), chunk_size):
        chunk = window.iloc[start:start + chunk_size]
        if chunk.empty:
            continue

        chunk_start = bounded_offset + start + 1
        chunk_end = chunk_start + len(chunk.index) - 1
        lines = ["\t".join(header)] if any(header) else []
        for row in chunk.itertuples(index=False, name=None):
            values = [stringify_value(value) for value in row]
            while values and not values[-1]:
                values.pop()
            lines.append("\t".join(values))

        blocks.append(
            {
                "kind": "rows",
                "sourceLabel": f"{source_name} rows {chunk_start}-{chunk_end}",
                "rowStart": chunk_start,
                "rowEnd": chunk_end,
                "text": normalize_text("\n".join(lines)),
            }
        )

    return blocks, total_rows, len(window.index)


def column_summary(column_name: str, series, max_top_values: int = 5) -> dict[str, Any]:
    import pandas as pd  # pylint: disable=import-outside-toplevel

    non_null = series.dropna()
    count = int(series.shape[0])
    non_null_count = int(non_null.shape[0])
    null_count = count - non_null_count
    unique_count = int(non_null.nunique(dropna=True)) if non_null_count else 0

    result: dict[str, Any] = {
        "name": str(column_name),
        "nonNullCount": non_null_count,
        "nullCount": null_count,
        "uniqueCount": unique_count,
    }

    if non_null_count == 0:
        result["kind"] = "empty"
        result["sampleValues"] = []
        return result

    numeric = pd.to_numeric(non_null, errors="coerce")
    numeric_ratio = float(numeric.notna().mean())
    if numeric_ratio >= 0.95:
        result["kind"] = "numeric"
        result["sampleValues"] = [stringify_value(value) for value in non_null.head(3).tolist()]
        numeric = numeric.dropna()
        result["summary"] = {
            "min": float(numeric.min()),
            "max": float(numeric.max()),
            "mean": float(numeric.mean()),
            "median": float(numeric.median()),
        }
        return result

    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        datetime_series = pd.to_datetime(non_null, errors="coerce")
    datetime_ratio = float(datetime_series.notna().mean())
    if datetime_ratio >= 0.95:
        valid_dates = datetime_series.dropna()
        result["kind"] = "datetime"
        result["sampleValues"] = [stringify_value(value) for value in non_null.head(3).tolist()]
        result["summary"] = {
            "min": valid_dates.min().isoformat(),
            "max": valid_dates.max().isoformat(),
        }
        return result

    normalized = non_null.astype(str).map(str.strip)
    sample_values = [value for value in normalized.head(3).tolist() if value]
    top_values = normalized.value_counts().head(max_top_values)
    result["sampleValues"] = sample_values
    result["topValues"] = [
        {"value": str(index), "count": int(value)}
        for index, value in top_values.items()
    ]
    result["kind"] = "categorical" if unique_count <= max(20, min(100, non_null_count // 5 or 1)) else "text"
    return result


def frame_row_records(frame, rows: int) -> list[dict[str, str]]:
    records: list[dict[str, str]] = []
    if rows <= 0 or frame.empty:
        return records
    for _, row in frame.head(rows).iterrows():
        records.append(
            {
                str(column): stringify_value(value)
                for column, value in row.items()
            }
        )
    return records


def analyze_tabular_frame(
    path: Path,
    workspace_root: Path,
    *,
    sheet: str | None = None,
    max_columns: int = 40,
    sample_rows: int = 5,
    focus_columns: list[str] | None = None,
) -> dict[str, Any]:
    frame, source_name, loader_metadata = load_tabular_frame(path, sheet=sheet)
    frame = frame.copy()
    total_rows = int(len(frame.index))
    total_columns = int(len(frame.columns))
    all_columns = [str(column) for column in frame.columns.tolist()]

    selected_columns = all_columns
    if focus_columns:
        selected_columns = [column for column in all_columns if column in focus_columns]
    truncated_columns = False
    if len(selected_columns) > max_columns:
        selected_columns = selected_columns[:max_columns]
        truncated_columns = True

    column_summaries = [
        column_summary(column, frame[column])
        for column in selected_columns
    ]

    numeric_columns = [item["name"] for item in column_summaries if item.get("kind") == "numeric"]
    categorical_columns = [item["name"] for item in column_summaries if item.get("kind") == "categorical"]
    datetime_columns = [item["name"] for item in column_summaries if item.get("kind") == "datetime"]

    metadata = {
        "rowCount": total_rows,
        "columnCount": total_columns,
        "columns": all_columns,
        "analyzedColumns": selected_columns,
        "numericColumns": numeric_columns,
        "categoricalColumns": categorical_columns,
        "datetimeColumns": datetime_columns,
        "truncatedColumns": truncated_columns,
        "sampleRowCount": sample_rows,
    }
    metadata.update(loader_metadata)

    return {
        "path": relative_path(path, workspace_root),
        "documentType": "table",
        "title": path.stem,
        "sourceName": source_name,
        "metadata": metadata,
        "columns": column_summaries,
        "sampleRows": {
            "head": frame_row_records(frame.head(sample_rows), sample_rows),
            "tail": frame_row_records(frame.tail(sample_rows), sample_rows),
        },
    }


def normalize_logic(raw: Any) -> str:
    logic = stringify_value(raw or "and").lower()
    if logic not in {"and", "or"}:
        raise RuntimeError("filter_logic must be either 'and' or 'or'.")
    return logic


def resolve_column_name(column_name: Any, available_columns: list[str]) -> str:
    requested = stringify_value(column_name)
    if not requested:
        raise RuntimeError("Column name is required.")
    if requested in available_columns:
        return requested

    lowered = requested.lower()
    lowered_matches = [column for column in available_columns if column.lower() == lowered]
    if len(lowered_matches) == 1:
        return lowered_matches[0]
    if len(lowered_matches) > 1:
        raise RuntimeError(
            f"Column '{requested}' is ambiguous. Please use the exact column name."
        )

    preview = ", ".join(available_columns[:20])
    suffix = "..." if len(available_columns) > 20 else ""
    raise RuntimeError(
        f"Column '{requested}' was not found. Available columns: {preview}{suffix}"
    )


def resolve_column_list(requested_columns: list[str], available_columns: list[str]) -> list[str]:
    if not requested_columns:
        return []
    resolved: list[str] = []
    seen: set[str] = set()
    for requested in requested_columns:
        if requested == "*":
            for column in available_columns:
                if column not in seen:
                    resolved.append(column)
                    seen.add(column)
            continue
        column = resolve_column_name(requested, available_columns)
        if column not in seen:
            resolved.append(column)
            seen.add(column)
    return resolved


def normalize_text_series(series):
    return series.map(stringify_value).astype(str).str.strip()


def try_parse_number(value: Any) -> float | None:
    if value is None or isinstance(value, bool):
        return None
    if isinstance(value, (int, float)):
        if isinstance(value, float) and value != value:
            return None
        return float(value)
    if isinstance(value, str):
        candidate = value.strip().replace(",", "")
        if not candidate:
            return None
        try:
            return float(candidate)
        except ValueError:
            return None
    return None


def try_parse_datetime_value(value: Any):
    try:
        import pandas as pd  # pylint: disable=import-outside-toplevel
    except Exception:  # pragma: no cover
        return None

    if value in (None, ""):
        return None

    if isinstance(value, (datetime, date)):
        return pd.Timestamp(value)

    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        parsed = pd.to_datetime(value, errors="coerce")
    if pd.isna(parsed):
        return None
    return parsed


def coerce_filter_values(raw_value: Any, raw_values: Any) -> list[Any]:
    if isinstance(raw_values, list):
        return raw_values
    if raw_values not in (None, ""):
        return [raw_values]
    if isinstance(raw_value, list):
        return raw_value
    return [raw_value]


def evaluate_scalar_filter(series, operator: str, raw_value: Any):
    import pandas as pd  # pylint: disable=import-outside-toplevel

    if operator in {"is_empty", "not_empty"}:
        text_series = normalize_text_series(series)
        mask = text_series.eq("")
        return mask if operator == "is_empty" else ~mask

    if operator in {"contains", "not_contains", "starts_with", "ends_with"}:
        needle = stringify_value(raw_value)
        if not needle:
            raise RuntimeError(f"Filter operator '{operator}' requires a non-empty value.")
        text_series = normalize_text_series(series).str.casefold()
        needle_cf = needle.casefold()
        if operator == "contains":
            mask = text_series.str.contains(needle_cf, regex=False, na=False)
        elif operator == "not_contains":
            mask = ~text_series.str.contains(needle_cf, regex=False, na=False)
        elif operator == "starts_with":
            mask = text_series.str.startswith(needle_cf, na=False)
        else:
            mask = text_series.str.endswith(needle_cf, na=False)
        return mask

    text_series = normalize_text_series(series).str.casefold()
    value_text = stringify_value(raw_value).casefold()

    numeric_value = try_parse_number(raw_value)
    if numeric_value is not None:
        numeric_series = pd.to_numeric(series, errors="coerce")
        populated = max(1, int(series.notna().sum()))
        if int(numeric_series.notna().sum()) >= max(1, populated // 2):
            if operator in {"=", "==", "eq"}:
                return numeric_series.eq(numeric_value).fillna(False)
            if operator in {"!=", "<>", "neq"}:
                return numeric_series.ne(numeric_value).fillna(False)
            if operator == ">":
                return numeric_series.gt(numeric_value).fillna(False)
            if operator == ">=":
                return numeric_series.ge(numeric_value).fillna(False)
            if operator == "<":
                return numeric_series.lt(numeric_value).fillna(False)
            if operator == "<=":
                return numeric_series.le(numeric_value).fillna(False)

    datetime_value = try_parse_datetime_value(raw_value)
    if datetime_value is not None:
        with warnings.catch_warnings():
            warnings.simplefilter("ignore")
            datetime_series = pd.to_datetime(series, errors="coerce")
        populated = max(1, int(series.notna().sum()))
        if int(datetime_series.notna().sum()) >= max(1, populated // 2):
            if operator in {"=", "==", "eq"}:
                return datetime_series.eq(datetime_value).fillna(False)
            if operator in {"!=", "<>", "neq"}:
                return datetime_series.ne(datetime_value).fillna(False)
            if operator == ">":
                return datetime_series.gt(datetime_value).fillna(False)
            if operator == ">=":
                return datetime_series.ge(datetime_value).fillna(False)
            if operator == "<":
                return datetime_series.lt(datetime_value).fillna(False)
            if operator == "<=":
                return datetime_series.le(datetime_value).fillna(False)

    if operator in {"=", "==", "eq"}:
        return text_series.eq(value_text)
    if operator in {"!=", "<>", "neq"}:
        return text_series.ne(value_text)
    if operator == ">":
        return text_series.gt(value_text)
    if operator == ">=":
        return text_series.ge(value_text)
    if operator == "<":
        return text_series.lt(value_text)
    if operator == "<=":
        return text_series.le(value_text)

    raise RuntimeError(f"Unsupported filter operator: {operator}")


def evaluate_filter(series, spec: dict[str, Any]):
    import pandas as pd  # pylint: disable=import-outside-toplevel

    operator = stringify_value(spec.get("op") or spec.get("operator")).lower()
    if not operator:
        raise RuntimeError("Each filter must include an 'op'.")

    if operator in {"in", "not_in"}:
        values = coerce_filter_values(spec.get("value"), spec.get("values"))
        if not values:
            raise RuntimeError(f"Filter operator '{operator}' requires 'values'.")

        numeric_values = [try_parse_number(value) for value in values]
        if all(value is not None for value in numeric_values):
            numeric_series = pd.to_numeric(series, errors="coerce")
            populated = max(1, int(series.notna().sum()))
            if int(numeric_series.notna().sum()) >= max(1, populated // 2):
                mask = numeric_series.isin([value for value in numeric_values if value is not None]).fillna(False)
                return mask if operator == "in" else ~mask

        datetime_values = [try_parse_datetime_value(value) for value in values]
        if all(value is not None for value in datetime_values):
            with warnings.catch_warnings():
                warnings.simplefilter("ignore")
                datetime_series = pd.to_datetime(series, errors="coerce")
            populated = max(1, int(series.notna().sum()))
            if int(datetime_series.notna().sum()) >= max(1, populated // 2):
                normalized_values = [value for value in datetime_values if value is not None]
                mask = datetime_series.isin(normalized_values).fillna(False)
                return mask if operator == "in" else ~mask

        text_values = {
            stringify_value(value).casefold()
            for value in values
            if stringify_value(value)
        }
        text_series = normalize_text_series(series).str.casefold()
        mask = text_series.isin(text_values)
        return mask if operator == "in" else ~mask

    return evaluate_scalar_filter(series, operator, spec.get("value"))


def parse_sort_specs(raw: Any) -> list[dict[str, str]]:
    if raw in (None, "", []):
        return []

    if isinstance(raw, str) and not raw.strip().startswith("["):
        specs: list[dict[str, str]] = []
        for item in raw.split(","):
            token = item.strip()
            if not token:
                continue
            parts = [part.strip() for part in token.split(":") if part.strip()]
            if len(parts) == 1:
                specs.append({"column": parts[0], "direction": "asc"})
            else:
                specs.append({"column": parts[0], "direction": parts[1]})
        return specs

    return parse_json_items(raw, "sort_by")


def apply_sort(frame, raw_sort_specs: list[dict[str, Any]]):
    if not raw_sort_specs or frame.empty:
        return frame

    available_columns = [str(column) for column in frame.columns.tolist()]
    sort_columns: list[str] = []
    ascending: list[bool] = []
    for index, spec in enumerate(raw_sort_specs, start=1):
        column = resolve_column_name(spec.get("column"), available_columns)
        direction = stringify_value(spec.get("direction") or "asc").lower()
        if direction not in {"asc", "desc"}:
            raise RuntimeError(f"sort_by[{index}] direction must be 'asc' or 'desc'.")
        sort_columns.append(column)
        ascending.append(direction == "asc")

    return frame.sort_values(
        by=sort_columns,
        ascending=ascending,
        na_position="last",
        kind="mergesort",
    )


def normalize_aggregation_op(raw: Any) -> str:
    op = stringify_value(raw).lower()
    if op == "avg":
        return "mean"
    return op


def default_aggregation_alias(op: str, column: str | None) -> str:
    if op == "count" and column in (None, "", "*"):
        return "row_count"
    suffix = "rows" if column in (None, "", "*") else str(column)
    return f"{op}_{suffix}"


def prepare_aggregation_specs(
    raw_aggregations: list[dict[str, Any]],
    available_columns: list[str],
) -> list[dict[str, Any]]:
    specs: list[dict[str, Any]] = []
    seen_aliases: set[str] = set()

    for index, item in enumerate(raw_aggregations, start=1):
        op = normalize_aggregation_op(item.get("op"))
        if op not in {"count", "sum", "mean", "min", "max", "nunique"}:
            raise RuntimeError(
                "aggregations only support count, sum, avg/mean, min, max, nunique."
            )

        raw_column = item.get("column")
        column = None if raw_column in (None, "", "*") else resolve_column_name(raw_column, available_columns)
        alias = stringify_value(item.get("as")) or default_aggregation_alias(op, column or stringify_value(raw_column))
        if not alias:
            raise RuntimeError(f"aggregations[{index}] is missing a valid alias.")
        if alias in seen_aliases:
            raise RuntimeError(f"Duplicate aggregation alias '{alias}'.")
        seen_aliases.add(alias)

        specs.append(
            {
                "alias": alias,
                "column": column,
                "op": op,
                "rowCount": op == "count" and raw_column in (None, "", "*"),
            }
        )
    return specs


def compute_aggregation_value(frame, spec: dict[str, Any]):
    if spec["rowCount"]:
        return len(frame.index)

    column = spec["column"]
    series = frame[column]
    op = spec["op"]

    if op == "count":
        return int(series.count())
    if op == "nunique":
        return int(series.nunique(dropna=True))
    if op in {"sum", "mean"}:
        try:
            import pandas as pd  # pylint: disable=import-outside-toplevel
        except Exception as exc:  # pragma: no cover
            raise RuntimeError(f"pandas is required for aggregation: {exc}") from exc
        numeric = pd.to_numeric(series, errors="coerce")
        if op == "sum":
            return numeric.sum(min_count=1)
        return numeric.mean()
    if op == "min":
        return series.min()
    if op == "max":
        return series.max()

    raise RuntimeError(f"Unsupported aggregation op: {op}")


def build_aggregation_frame(frame, group_by_columns: list[str], aggregation_specs: list[dict[str, Any]]):
    import pandas as pd  # pylint: disable=import-outside-toplevel

    if not aggregation_specs:
        aggregation_specs = [{"alias": "row_count", "column": None, "op": "count", "rowCount": True}]

    if group_by_columns:
        grouped = frame.groupby(group_by_columns, dropna=False, sort=False)
        base = grouped.size().reset_index(name="__group_size__")
        result = base[group_by_columns].copy()
        for spec in aggregation_specs:
            if spec["rowCount"]:
                result[spec["alias"]] = base["__group_size__"].tolist()
            else:
                grouped_values = grouped[spec["column"]]
                if spec["op"] == "count":
                    values = grouped_values.count().reset_index(name=spec["alias"])[spec["alias"]]
                elif spec["op"] == "nunique":
                    values = grouped_values.nunique(dropna=True).reset_index(name=spec["alias"])[spec["alias"]]
                elif spec["op"] in {"sum", "mean"}:
                    numeric_series = pd.to_numeric(frame[spec["column"]], errors="coerce")
                    numeric_grouped = numeric_series.groupby(
                        [frame[column] for column in group_by_columns],
                        dropna=False,
                        sort=False,
                    )
                    if spec["op"] == "sum":
                        values = numeric_grouped.sum(min_count=1).reset_index(name=spec["alias"])[spec["alias"]]
                    else:
                        values = numeric_grouped.mean().reset_index(name=spec["alias"])[spec["alias"]]
                elif spec["op"] == "min":
                    values = grouped_values.min().reset_index(name=spec["alias"])[spec["alias"]]
                elif spec["op"] == "max":
                    values = grouped_values.max().reset_index(name=spec["alias"])[spec["alias"]]
                else:
                    raise RuntimeError(f"Unsupported aggregation op: {spec['op']}")
                result[spec["alias"]] = values.tolist()
        return result

    row = {
        spec["alias"]: compute_aggregation_value(frame, spec)
        for spec in aggregation_specs
    }
    return pd.DataFrame([row])


def to_json_compatible(value: Any) -> Any:
    try:
        import pandas as pd  # pylint: disable=import-outside-toplevel
    except Exception:  # pragma: no cover
        pd = None  # type: ignore[assignment]

    if value is None:
        return None
    if hasattr(value, "item") and not isinstance(value, (str, bytes)):
        try:
            value = value.item()
        except Exception:
            pass
    if isinstance(value, (datetime, date)):
        return value.isoformat()
    if pd is not None:
        try:
            if pd.isna(value):
                return None
        except Exception:
            pass
        if isinstance(value, pd.Timestamp):
            return value.isoformat()
    if isinstance(value, float) and value != value:
        return None
    if isinstance(value, (bool, int, float, str)):
        return value
    return str(value)


def frame_to_query_rows(frame) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for record in frame.to_dict(orient="records"):
        rows.append(
            {
                str(column): to_json_compatible(value)
                for column, value in record.items()
            }
        )
    return rows


def query_tabular_frame(
    path: Path,
    workspace_root: Path,
    *,
    sheet: str | None = None,
    select_columns: list[str] | None = None,
    filters: list[dict[str, Any]] | None = None,
    filter_logic: str = "and",
    group_by: list[str] | None = None,
    aggregations: list[dict[str, Any]] | None = None,
    sort_by: list[dict[str, Any]] | None = None,
    row_offset: int = 0,
    limit: int = 50,
) -> dict[str, Any]:
    import pandas as pd  # pylint: disable=import-outside-toplevel

    frame, source_name, loader_metadata = load_tabular_frame(path, sheet=sheet)
    frame = frame.copy()
    all_columns = [str(column) for column in frame.columns.tolist()]
    total_rows = int(len(frame.index))

    requested_select_columns = resolve_column_list(select_columns or [], all_columns)
    group_by_columns = resolve_column_list(group_by or [], all_columns)
    aggregation_specs = prepare_aggregation_specs(aggregations or [], all_columns)
    sort_specs = sort_by or []

    working = frame
    if filters:
        logic = normalize_logic(filter_logic)
        masks = []
        for index, item in enumerate(filters, start=1):
            column = resolve_column_name(item.get("column"), all_columns)
            spec = dict(item)
            spec["column"] = column
            try:
                masks.append(evaluate_filter(working[column], spec).fillna(False))
            except Exception as exc:
                raise RuntimeError(f"Invalid filter at position {index}: {exc}") from exc
        if masks:
            if logic == "and":
                combined = masks[0]
                for mask in masks[1:]:
                    combined = combined & mask
            else:
                combined = masks[0]
                for mask in masks[1:]:
                    combined = combined | mask
            working = working.loc[combined].copy()

    matched_row_count = int(len(working.index))
    query_mode = "rows"

    if group_by_columns or aggregation_specs:
        query_mode = "grouped_aggregate" if group_by_columns else "aggregate"
        result_frame = build_aggregation_frame(working, group_by_columns, aggregation_specs)
        result_frame = apply_sort(result_frame, sort_specs)
        available_result_columns = [str(column) for column in result_frame.columns.tolist()]
        selected_result_columns = resolve_column_list(requested_select_columns, available_result_columns) or available_result_columns
        result_frame = result_frame[selected_result_columns]
    else:
        sorted_frame = apply_sort(working, sort_specs)
        selected_columns = requested_select_columns or all_columns
        result_frame = sorted_frame[selected_columns]

    total_result_rows = int(len(result_frame.index))
    bounded_offset = max(0, min(row_offset, total_result_rows))
    paged_frame = result_frame.iloc[bounded_offset:bounded_offset + limit].copy()
    returned_rows = int(len(paged_frame.index))
    selected_output_columns = [str(column) for column in paged_frame.columns.tolist()]

    metadata = {
        "rowCount": total_rows,
        "matchedRowCount": matched_row_count,
        "totalResultRowCount": total_result_rows,
        "returnedRowCount": returned_rows,
        "columnCount": len(all_columns),
        "columns": all_columns,
        "selectedColumns": selected_output_columns,
        "rowOffset": bounded_offset,
        "limit": limit,
        "truncated": bounded_offset + returned_rows < total_result_rows,
        "queryMode": query_mode,
        "query": {
            "filterCount": len(filters or []),
            "filterLogic": normalize_logic(filter_logic),
            "groupBy": group_by_columns,
            "aggregationCount": len(aggregation_specs),
            "sortCount": len(sort_specs),
        },
    }
    metadata.update(loader_metadata)

    return {
        "path": relative_path(path, workspace_root),
        "documentType": "table_query",
        "title": path.stem,
        "sourceName": source_name,
        "metadata": metadata,
        "rows": frame_to_query_rows(paged_frame),
    }


def extract_pdf(
    path: Path,
    workspace_root: Path,
    max_chars: int,
    max_blocks: int,
) -> dict[str, Any]:
    reader = PdfReader(str(path))
    blocks: list[dict[str, Any]] = []
    for page_index, page in enumerate(reader.pages, start=1):
        text = normalize_text(page.extract_text() or "")
        if not text:
            continue
        blocks.append(
            {
                "kind": "page",
                "sourceLabel": f"Page {page_index}",
                "page": page_index,
                "text": text,
            }
        )

    metadata = {
        "pageCount": len(reader.pages),
    }
    title = None
    if reader.metadata and getattr(reader.metadata, "title", None):
        candidate = str(reader.metadata.title).strip()
        title = candidate if candidate and candidate.lower() != "untitled" else None
    return build_result(
        path=path,
        workspace_root=workspace_root,
        document_type="pdf",
        title=title,
        metadata=metadata,
        blocks=blocks,
        max_chars=max_chars,
        max_blocks=max_blocks,
    )


def extract_docx(
    path: Path,
    workspace_root: Path,
    max_chars: int,
    max_blocks: int,
) -> dict[str, Any]:
    doc = Document(str(path))
    blocks: list[dict[str, Any]] = []

    for index, paragraph in enumerate(doc.paragraphs, start=1):
        text = normalize_text(paragraph.text)
        if not text:
            continue
        blocks.append(
            {
                "kind": "paragraph",
                "sourceLabel": f"Paragraph {index}",
                "paragraph": index,
                "text": text,
            }
        )

    for table_index, table in enumerate(doc.tables, start=1):
        lines: list[str] = []
        for row in table.rows:
            values = [stringify_value(cell.text) for cell in row.cells]
            trimmed = list(values)
            while trimmed and not trimmed[-1]:
                trimmed.pop()
            if trimmed:
                lines.append("\t".join(trimmed))
        text = normalize_text("\n".join(lines))
        if not text:
            continue
        blocks.append(
            {
                "kind": "table",
                "sourceLabel": f"Table {table_index}",
                "table": table_index,
                "text": text,
            }
        )

    core_title = getattr(doc.core_properties, "title", None)
    metadata = {
        "paragraphCount": len(doc.paragraphs),
        "tableCount": len(doc.tables),
    }
    return build_result(
        path=path,
        workspace_root=workspace_root,
        document_type="docx",
        title=str(core_title) if core_title else None,
        metadata=metadata,
        blocks=blocks,
        max_chars=max_chars,
        max_blocks=max_blocks,
    )


def rows_to_sheet_blocks(
    sheet_name: str,
    rows: list[tuple[int, list[str]]],
    chunk_size: int = 40,
) -> list[dict[str, Any]]:
    blocks: list[dict[str, Any]] = []
    for start in range(0, len(rows), chunk_size):
        chunk = rows[start:start + chunk_size]
        if not chunk:
            continue
        lines = ["\t".join(values) for _, values in chunk]
        max_width = max(len(values) for _, values in chunk)
        first_row = chunk[0][0]
        last_row = chunk[-1][0]
        cell_range = f"A{first_row}:{get_column_letter(max_width)}{last_row}"
        blocks.append(
            {
                "kind": "sheet",
                "sourceLabel": f"{sheet_name}!{cell_range}",
                "sheet": sheet_name,
                "cellRange": cell_range,
                "rowStart": first_row,
                "rowEnd": last_row,
                "text": normalize_text("\n".join(lines)),
            }
        )
    return blocks


def extract_xlsx(
    path: Path,
    workspace_root: Path,
    max_chars: int,
    max_blocks: int,
    *,
    sheet: str | None = None,
    row_offset: int = 0,
    max_rows: int | None = None,
) -> dict[str, Any]:
    if sheet or row_offset or max_rows:
        frame, source_name, loader_metadata = load_tabular_frame(path, sheet=sheet)
        blocks, total_rows, returned_rows = frame_window_to_blocks(
            source_name,
            frame,
            row_offset=row_offset,
            max_rows=max_rows,
        )
        metadata = {
            "sheetCount": len(loader_metadata.get("sheetNames", [source_name])),
            "sheetNames": loader_metadata.get("sheetNames", [source_name]),
            "sheet": loader_metadata.get("sheetName", source_name),
            "rowCount": total_rows,
            "returnedRowCount": returned_rows,
            "rowOffset": row_offset,
            "columnCount": int(len(frame.columns)),
            "columns": [str(column) for column in frame.columns.tolist()],
        }
    else:
        workbook = load_workbook(filename=str(path), read_only=True, data_only=True)
        blocks = []
        sheet_names: list[str] = []
        sheet_summaries: list[dict[str, Any]] = []

        for worksheet in workbook.worksheets:
            sheet_names.append(worksheet.title)
            rows: list[tuple[int, list[str]]] = []
            max_width = 0
            for row_index, row in enumerate(worksheet.iter_rows(values_only=True), start=1):
                values = [stringify_value(value) for value in row]
                while values and not values[-1]:
                    values.pop()
                if not values:
                    continue
                max_width = max(max_width, len(values))
                rows.append((row_index, values))

            blocks.extend(rows_to_sheet_blocks(worksheet.title, rows))
            sheet_summaries.append(
                {
                    "sheet": worksheet.title,
                    "rowCount": len(rows),
                    "columnCount": max_width,
                    "dimensions": worksheet.calculate_dimension(),
                }
            )

        metadata = {
            "sheetCount": len(sheet_names),
            "sheetNames": sheet_names,
            "sheets": sheet_summaries,
        }
    return build_result(
        path=path,
        workspace_root=workspace_root,
        document_type="spreadsheet",
        title=path.stem,
        metadata=metadata,
        blocks=blocks,
        max_chars=max_chars,
        max_blocks=max_blocks,
    )


def extract_legacy_xls(
    path: Path,
    workspace_root: Path,
    max_chars: int,
    max_blocks: int,
) -> dict[str, Any]:
    try:
        import pandas as pd  # pylint: disable=import-outside-toplevel
    except Exception as exc:  # pragma: no cover - import failures are surfaced to user
        raise RuntimeError(f"pandas is required to read legacy .xls files: {exc}") from exc

    try:
        sheets = pd.read_excel(str(path), sheet_name=None, dtype=str)
    except Exception as exc:
        raise RuntimeError(
            "Legacy .xls files are not supported by the current runtime. "
            "Please convert the file to .xlsx or install an .xls reader such as xlrd."
        ) from exc

    blocks: list[dict[str, Any]] = []
    sheet_names: list[str] = []
    sheet_summaries: list[dict[str, Any]] = []

    for sheet_name, frame in sheets.items():
        sheet_names.append(str(sheet_name))
        shaped = frame.fillna("").astype(str)
        rows: list[tuple[int, list[str]]] = []
        header = [stringify_value(column) for column in shaped.columns.tolist()]
        if any(header):
            rows.append((1, header))
        for data_index, row in enumerate(shaped.itertuples(index=False, name=None), start=2):
            values = [stringify_value(value) for value in row]
            while values and not values[-1]:
                values.pop()
            if values:
                rows.append((data_index, values))
        blocks.extend(rows_to_sheet_blocks(str(sheet_name), rows))
        sheet_summaries.append(
            {
                "sheet": str(sheet_name),
                "rowCount": len(shaped.index),
                "columnCount": len(shaped.columns),
            }
        )

    metadata = {
        "sheetCount": len(sheet_names),
        "sheetNames": sheet_names,
        "sheets": sheet_summaries,
    }
    return build_result(
        path=path,
        workspace_root=workspace_root,
        document_type="spreadsheet",
        title=path.stem,
        metadata=metadata,
        blocks=blocks,
        max_chars=max_chars,
        max_blocks=max_blocks,
    )


def extract_delimited_table(
    path: Path,
    workspace_root: Path,
    max_chars: int,
    max_blocks: int,
    *,
    row_offset: int = 0,
    max_rows: int | None = None,
) -> dict[str, Any]:
    frame, source_name, _ = load_tabular_frame(path)
    blocks, total_rows, returned_rows = frame_window_to_blocks(
        source_name,
        frame,
        row_offset=row_offset,
        max_rows=max_rows,
    )
    metadata = {
        "rowCount": total_rows,
        "returnedRowCount": returned_rows,
        "rowOffset": row_offset,
        "columnCount": len(frame.columns),
        "columns": [str(column) for column in frame.columns.tolist()],
    }
    return build_result(
        path=path,
        workspace_root=workspace_root,
        document_type="table",
        title=path.stem,
        metadata=metadata,
        blocks=blocks,
        max_chars=max_chars,
        max_blocks=max_blocks,
    )


def extract_text_document(
    path: Path,
    workspace_root: Path,
    max_chars: int,
    max_blocks: int,
) -> dict[str, Any]:
    text = path.read_text(encoding="utf-8", errors="replace")
    segments = [normalize_text(segment) for segment in text.split("\n\n")]
    blocks = []
    for index, segment in enumerate(segments, start=1):
        if not segment:
            continue
        blocks.append(
            {
                "kind": "section",
                "sourceLabel": f"Section {index}",
                "text": segment,
            }
        )

    metadata = {
        "extension": path.suffix.lower(),
    }
    return build_result(
        path=path,
        workspace_root=workspace_root,
        document_type="text",
        title=path.stem,
        metadata=metadata,
        blocks=blocks,
        max_chars=max_chars,
        max_blocks=max_blocks,
    )


def extract_html_document(
    path: Path,
    workspace_root: Path,
    max_chars: int,
    max_blocks: int,
) -> dict[str, Any]:
    raw = path.read_text(encoding="utf-8", errors="replace")
    parser = HtmlTextExtractor()
    parser.feed(raw)
    text = parser.text()
    segments = [normalize_text(segment) for segment in text.split("\n\n")]
    blocks = []
    for index, segment in enumerate(segments, start=1):
        if not segment:
            continue
        blocks.append(
            {
                "kind": "html-section",
                "sourceLabel": f"Section {index}",
                "text": segment,
            }
        )

    metadata = {
        "extension": path.suffix.lower(),
    }
    return build_result(
        path=path,
        workspace_root=workspace_root,
        document_type="html",
        title=path.stem,
        metadata=metadata,
        blocks=blocks,
        max_chars=max_chars,
        max_blocks=max_blocks,
    )


def extract_document(
    path: Path,
    workspace_root: Path,
    max_chars: int,
    max_blocks: int,
    *,
    sheet: str | None = None,
    row_offset: int = 0,
    max_rows: int | None = None,
) -> dict[str, Any]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(path, workspace_root, max_chars, max_blocks)
    if suffix == ".docx":
        return extract_docx(path, workspace_root, max_chars, max_blocks)
    if suffix == ".xlsx":
        return extract_xlsx(
            path,
            workspace_root,
            max_chars,
            max_blocks,
            sheet=sheet,
            row_offset=row_offset,
            max_rows=max_rows,
        )
    if suffix == ".xls":
        return extract_legacy_xls(path, workspace_root, max_chars, max_blocks)
    if suffix in {".csv", ".tsv"}:
        return extract_delimited_table(
            path,
            workspace_root,
            max_chars,
            max_blocks,
            row_offset=row_offset,
            max_rows=max_rows,
        )
    if suffix in TEXT_EXTENSIONS:
        return extract_text_document(path, workspace_root, max_chars, max_blocks)
    if suffix in HTML_EXTENSIONS:
        return extract_html_document(path, workspace_root, max_chars, max_blocks)
    raise RuntimeError(f"Unsupported document type: {suffix or '[no extension]'}")


def handle_read_document(payload: dict[str, Any]) -> dict[str, Any]:
    workspace_root = Path(payload["workspaceRoot"]).resolve()
    path = Path(payload["path"]).resolve()
    max_chars = int(payload.get("maxChars") or 6000)
    max_blocks = int(payload.get("maxBlocks") or 24)
    sheet = payload.get("sheet")
    row_offset = parse_optional_non_negative_int(payload.get("rowOffset"), default=0)
    max_rows = parse_optional_positive_int(payload.get("maxRows"))
    return extract_document(
        path,
        workspace_root,
        max_chars,
        max_blocks,
        sheet=sheet if isinstance(sheet, str) and sheet.strip() else None,
        row_offset=row_offset,
        max_rows=max_rows,
    )


def handle_analyze_tabular_document(payload: dict[str, Any]) -> dict[str, Any]:
    workspace_root = Path(payload["workspaceRoot"]).resolve()
    path = Path(payload["path"]).resolve()
    sheet = payload.get("sheet")
    max_columns = int(payload.get("maxColumns") or 40)
    sample_rows = int(payload.get("sampleRows") or 5)
    focus_columns = parse_focus_columns(payload.get("focusColumns"))
    return analyze_tabular_frame(
        path,
        workspace_root,
        sheet=sheet if isinstance(sheet, str) and sheet.strip() else None,
        max_columns=max(1, min(max_columns, 200)),
        sample_rows=max(1, min(sample_rows, 20)),
        focus_columns=focus_columns or None,
    )


def handle_query_tabular_document(payload: dict[str, Any]) -> dict[str, Any]:
    workspace_root = Path(payload["workspaceRoot"]).resolve()
    path = Path(payload["path"]).resolve()
    sheet = payload.get("sheet")
    return query_tabular_frame(
        path,
        workspace_root,
        sheet=sheet if isinstance(sheet, str) and sheet.strip() else None,
        select_columns=parse_comma_separated(payload.get("selectColumns")) or None,
        filters=parse_json_items(payload.get("filters"), "filters") or None,
        filter_logic=payload.get("filterLogic") or "and",
        group_by=parse_comma_separated(payload.get("groupBy")) or None,
        aggregations=parse_json_items(payload.get("aggregations"), "aggregations") or None,
        sort_by=parse_sort_specs(payload.get("sortBy")) or None,
        row_offset=parse_optional_non_negative_int(payload.get("rowOffset"), default=0),
        limit=max(1, min(int(payload.get("limit") or 50), 500)),
    )


def iter_document_paths(root_path: Path, extensions: set[str]) -> list[Path]:
    matches: list[Path] = []
    for current_root, dirs, files in os.walk(root_path):
        dirs[:] = [
            directory
            for directory in sorted(dirs)
            if directory not in SKIP_DIRS and not directory.startswith(".")
        ]
        for name in sorted(files):
            if name.startswith(".") or name.startswith("~$"):
                continue
            candidate = Path(current_root) / name
            if candidate.suffix.lower() in extensions:
                matches.append(candidate)
    return matches


def handle_index_workspace_documents(payload: dict[str, Any]) -> dict[str, Any]:
    workspace_root = Path(payload["workspaceRoot"]).resolve()
    root_path = Path(payload["path"]).resolve()
    max_files = max(1, int(payload.get("maxFiles") or 8))
    max_chars_per_file = max(200, int(payload.get("maxCharsPerFile") or 700))
    extensions = parse_extensions(payload.get("extensions"))

    matched_paths = iter_document_paths(root_path, extensions)
    selected_paths = matched_paths[:max_files]
    files: list[dict[str, Any]] = []
    skipped: list[dict[str, str]] = []

    for document_path in selected_paths:
        try:
            extracted = extract_document(
                document_path,
                workspace_root,
                max_chars=max_chars_per_file,
                max_blocks=6,
            )
            files.append(
                {
                    "path": extracted["path"],
                    "documentType": extracted["documentType"],
                    "title": extracted.get("title"),
                    "preview": extracted.get("content", ""),
                    "charCount": extracted.get("charCount", 0),
                    "truncated": extracted.get("truncated", False),
                    "blockCount": extracted.get("metadata", {}).get("blockCount", 0),
                    "metadata": extracted.get("metadata", {}),
                }
            )
        except Exception as exc:  # pragma: no cover - surfaced to UI/tool result
            skipped.append(
                {
                    "path": relative_path(document_path, workspace_root),
                    "reason": str(exc),
                }
            )

    return {
        "rootPath": relative_path(root_path, workspace_root),
        "supportedExtensions": sorted(extensions),
        "matchedFiles": len(matched_paths),
        "indexedFiles": len(files),
        "scanLimited": len(matched_paths) > len(selected_paths),
        "files": files,
        "skipped": skipped,
    }


def main() -> None:
    payload = json.loads(sys.stdin.read() or "{}")
    command = payload.get("command")
    if command == "read_document":
        result = handle_read_document(payload)
    elif command == "analyze_tabular_document":
        result = handle_analyze_tabular_document(payload)
    elif command == "query_tabular_document":
        result = handle_query_tabular_document(payload)
    elif command == "index_workspace_documents":
        result = handle_index_workspace_documents(payload)
    else:
        raise RuntimeError(f"Unsupported command: {command}")

    sys.stdout.write(json.dumps(result, ensure_ascii=False))


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:
        print(str(exc), file=sys.stderr)
        raise
